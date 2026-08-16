"""
Document extraction service.
Handles text extraction from PDF and DOCX files.
Follows Single Responsibility Principle — only does text extraction.
"""
import io
import logging
from typing import Tuple

import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)

# Supported file types
SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}

MAX_FILE_SIZE_MB = 10
MAX_CHARS = 50_000  # Limit context sent to Claude API


def validate_file(filename: str, content_type: str, file_size: int) -> Tuple[bool, str]:
    """
    Validates uploaded file before processing.
    Returns (is_valid, error_message).
    """
    # Check file size
    if file_size > MAX_FILE_SIZE_MB * 1024 * 1024:
        return False, f"File too large. Maximum size is {MAX_FILE_SIZE_MB}MB."

    # Check file type
    if content_type not in SUPPORTED_TYPES:
        supported = ", ".join([".pdf", ".docx", ".txt"])
        return False, f"Unsupported file type. Please upload: {supported}"

    return True, ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from a PDF file.
    Handles multi-page PDFs gracefully.
    """
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text.strip()}")
            except Exception as e:
                logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
                continue

        if not text_parts:
            raise ValueError("No readable text found in PDF. It may be a scanned image.")

        return "\n\n".join(text_parts)

    except PyPDF2.errors.PdfReadError as e:
        raise ValueError(f"Could not read PDF file: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a DOCX file.
    Preserves paragraph structure.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            raise ValueError("No readable text found in document.")

        return "\n\n".join(paragraphs)

    except Exception as e:
        raise ValueError(f"Could not read DOCX file: {str(e)}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    Decodes and returns plain text content.
    Tries UTF-8 first, falls back to latin-1.
    """
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1").strip()


def extract_text(file_bytes: bytes, content_type: str) -> str:
    """
    Main extraction dispatcher.
    Routes to the correct extractor based on file type.
    Truncates to MAX_CHARS to avoid API token limits.
    """
    file_type = SUPPORTED_TYPES.get(content_type, "unknown")

    if file_type == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        text = extract_text_from_docx(file_bytes)
    elif file_type == "txt":
        text = extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {content_type}")

    # Truncate if too long
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[Document truncated for processing...]"
        logger.info(f"Document truncated to {MAX_CHARS} characters")

    return text
